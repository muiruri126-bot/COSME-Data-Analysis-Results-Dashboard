from docx import Document
import os

base = r'C:\Users\Bmuiruri.PLANKE-KILIFI\Desktop\Benard\My Project\My work\Proposal'

# Document 1
print("=" * 80)
print("DOCUMENT 1: Project Design Reflection Questions response from Plan International Kenya")
print("=" * 80)
doc1 = Document(os.path.join(base, 'Project Design Reflection Questions  response from Plan International Kenya.docx'))
for p in doc1.paragraphs:
    print(p.text)

# Also extract tables
for i, table in enumerate(doc1.tables):
    print(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        print(" | ".join(cell.text for cell in row.cells))

print("\n\n")
print("=" * 80)
print("DOCUMENT 2: COSME Phase 2 Design Validation_Plan_International")
print("=" * 80)
doc2 = Document(os.path.join(base, 'COSME Phase 2 Design Validation_Plan_International.docx'))
for p in doc2.paragraphs:
    print(p.text)

for i, table in enumerate(doc2.tables):
    print(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        print(" | ".join(cell.text for cell in row.cells))
